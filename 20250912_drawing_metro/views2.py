# -*- coding: utf-8 -*-
"""
Created on Sun Apr 27 09:00:06 2025

@author: thinkpad
"""

import tkinter as tk
import tkinter.ttk as ttk
from tkinter import messagebox, filedialog
from tkinter import scrolledtext
import os
#from main import sorted_files
#from exportFigure1 import ex_Fig
import re
#import traceback
import threading
import queue
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image


class WordPage(ttk.Frame):
    def __init__(self, root):
        super().__init__(root)
        self.LoadFolder_path = tk.StringVar()   #输入图片地址
#        self.Open_PIC = tk.StringVar()  #输入设施地址
        self.SavePIC_path = tk.StringVar()  #输入导出word地址
        self.click_first = 0
        self.click_second = 0
        self.click_count = 0
        self.add_info = tk.BooleanVar()  #是否添加文字说明
        self.gap_info = tk.IntVar()  #间隔图片数量
        self.gap_info.set(1)
#        self.information = tk.StringVar()
#        self.information.set("每两站之间的运行时间、速度、网侧电流、电机电流、累计能耗、网侧输入功率、电机输出功率对运行里程及全线的速度、网侧电流、电机电流、累计能耗、网侧输入功率、电机输出功率对运行时间的曲线")
        #  后台运行状态
        self.is_running = False
        #  创建队列用于线程间通信
        self.queue = queue.Queue()
#        self.stop_event = threading.Event()
        self.after(100, self.process_queue)
        self.create_widget()
    def create_widget(self):
        self.group1 = ttk.Labelframe(self, text="设置图片打开路径和文档保存路径")
        self.group1.grid(row = 0, column = 0, padx = 10, pady = 10, rowspan=2)
        ttk.Label(self.group1, text="图片打开路径").grid(row=0, column=0, padx=10, pady=10)
        ttk.Entry(self.group1, textvariable = self.LoadFolder_path).grid(row=0, column=1, padx=10, pady=10)
        ttk.Button(self.group1, text="设置", command=self.open_folder).grid(row=0, column=2, padx=10, pady=10)
        
        ttk.Label(self.group1, text="文档保存路径").grid(row=1, column=0, padx=10, pady=10)
        ttk.Entry(self.group1, textvariable = self.SavePIC_path).grid(row=1, column=1, padx=10, pady=10)
        ttk.Button(self.group1, text="设置", command=self.save_folder).grid(row=1, column=2, padx=10, pady=10)
        
        self.listbox2 = tk.Listbox(self.group1, width=50 , selectmode=tk.MULTIPLE)
        self.listbox2.grid(row=2, column=0, padx=10, pady=10, columnspan=3)
        sc = ttk.Scrollbar(self.group1, command=self.listbox2.yview)
        sc.grid(row=2, column=4, sticky='ns')
        sc2 = ttk.Scrollbar(self.group1, orient='horizontal', command=self.listbox2.xview)
        sc2.grid(row=3, column=0, sticky='we', columnspan=3)
        # 定义右键多选功能
        self.listbox2.config(yscrollcommand=sc.set, xscrollcommand=sc2.set)
        self.listbox2.bind("<Button-3>", self.on_right_click)
#        sc.grid(row=0, column=0)
        # 创建删除按钮并绑定delete_selected函数
        add_button2 = ttk.Button(self.group1, text="添加", command = self.add_dir)
        add_button2.grid(row=4, column=0, padx=10, pady=10)
        delete_button2 = ttk.Button(self.group1, text="删除", command=self.delete_dir)
        delete_button2.grid(row=4, column=1, padx=10, pady=10)
        self.run_button2 = ttk.Button(self.group1, text="运行", command=self.create_progressbar)
        self.run_button2.grid(row=4, column=2, padx=10, pady=10)
        # 自定义信息输入
        self.group2 = ttk.Labelframe(self, text="设置自定义信息")
        self.group2.grid(row = 0, column = 1, padx = 10, pady = 10)
        ttk.Label(self.group2, text="是否增加自定义信息").grid(row=0, column=0, padx=10, pady=10, sticky="w")
        self.check_add = ttk.Checkbutton(self.group2, offvalue=False, onvalue=True, variable=self.add_info, command=self.toggle_state)
        self.check_add.grid(row=0, column=1, padx=10, pady=10)
        ttk.Label(self.group2, text = "间隔多少张图片写一段描述").grid(row=1, column=0, padx=10, pady=10, sticky="w")
        self.gap_entry = ttk.Entry(self.group2, width=3, textvariable=self.gap_info, state="disabled")
        self.gap_entry.grid(row=1, column=1, padx=10, pady=10)
        ttk.Label(self.group2, text = "文字描述内容（注：先检查图片命名格式是否为“序号 起点——终点 网压_编组_载荷_行车方向_曲线名称”）").grid(row=2, column=0, padx=10, pady=10)
        self.information_entry = scrolledtext.ScrolledText(self.group2, wrap=tk.WORD, width=80, height=10)
        self.information_entry.grid(row=3, column=0, padx=10, pady=10, columnspan=2, sticky="w")
        self.information_entry.config(state="disabled")
        
    def open_folder(self):
        try:
            # 打开文件夹选择对话框，并获取选择的文件夹路径
            folder_path = filedialog.askdirectory() #(initialdir = self.desktop_path)
            self.LoadFolder_path.set(folder_path)
            
        except Exception as e:
            messagebox.showwarning(title = '提示', message = '未选择任何文件夹')
    def save_folder(self):
        try:
            # 打开文件夹选择对话框，并获取选择的文件夹路径
            folder_path = filedialog.askdirectory() #(initialdir = self.desktop_path)
            self.SavePIC_path.set(folder_path)
            
        except Exception as e:
            messagebox.showwarning(title = '提示', message = '未选择任何文件夹')
#    def open_STA(self):
#        try:
#            filetypes = [("jpg", ".jpg"), ("jpeg", ".jpeg")]
#            filepath = filedialog.askopenfilename(title = '打开文件', 
#                      filetypes = filetypes, defaultextension = '.jpg')
#            self.Open_PIC.set(filepath)
#        except Exception as e:
#            messagebox.showwarning(title = '提示', message = '未选择任何图片')
        
    def add_dir(self):
        # 选择文件夹中的.jpg或.jpeg文件
        file_path = filedialog.askopenfilenames(initialdir=self.LoadFolder_path, title="选择图片", filetypes=(("jpg files", "*.jpg *.jpeg"),))
        if file_path:  # 如果用户选择了文件
            for fp in file_path:
                fp = os.path.basename(fp)
                self.listbox2.insert(tk.END, fp)
#            print(type(self.listbox2.get(0, tk.END)))
    def delete_dir(self):
        # 获取选中的项的索引列表
        selected_indices2 = self.listbox2.curselection()
        for index in sorted(selected_indices2, reverse=True):
            self.listbox2.delete(index)
            
    def on_right_click(self, event):
        # 获取当前鼠标点击的条目索引
        if self.click_count == 0:
            self.listbox2.selection_clear(0, tk.END)
            self.click_first = self.listbox2.nearest(event.y)
            self.listbox2.selection_set(self.click_first)
            self.click_count += 1
        else:
            self.click_second = self.listbox2.nearest(event.y)
            for i in range(min(self.click_first, self.click_second), max(self.click_first, self.click_second) + 1):
                self.listbox2.selection_set(i)
            self.click_count = 0
    def toggle_state(self):
    # 获取勾选框的状态
        if self.add_info.get():
            # 如果选中“是”，启用输入框和按钮
            self.gap_entry.config(state='normal')
            self.information_entry.config(state='normal')
            self.information_entry.insert(tk.END, "每两站之间的运行时间、速度、网侧电流、电机电流、累计能耗、网侧输入功率、电机输出功率、加速度对运行里程及全线的速度、网侧电流、电机电流、累计能耗、网侧输入功率、电机输出功率、加速度对运行时间的曲线")
        else:
            self.gap_entry.config(state='disabled')
            self.information_entry.delete('1.0', tk.END)
            self.information_entry.config(state='disabled')
    # 运行导入图片软件
    def create_progressbar(self):
        # 更新按钮状态
        self.run_button2.config(state=tk.DISABLED)        
        #创建进度条组件
        self.top = tk.Toplevel(self)
        self.top.title("任务进度")
        self.top.geometry("300x150")
        
        
#        self.task_length = 100
        self.progresslabel = tk.Label(self.top, text="进度:")
        self.progresslabel.pack(pady=10)
        # 进度条
#        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(
            self.top, 
            maximum=100,
            mode="determinate"
        )
        self.progress_bar.pack(pady=20, padx=20, fill=tk.X)
     
       
        # 停止按钮
        self.stop_button = tk.Button(
            self.top, 
            text="停止任务", 
            command=self.stop_task,
            state=tk.NORMAL
        )
        self.stop_button.pack(pady=10)

        # 重置停止事件
#        self.stop_event.clear()
        
        # 重置进度条
        self.progress_bar['value'] = 0

        # 启动任务线程
        self.thread = threading.Thread(target=self.run_program, daemon=True)
        self.thread.start()
        # 窗口关闭时的事件处理
        self.top.protocol("WM_DELETE_WINDOW", self.stop_task)
    def run_program(self):
        # 获取图片打开地址和文档存储地址
        Load = self.LoadFolder_path.get()
        Save = self.SavePIC_path.get()
        # 初始化图名数组
        file_names = []
        # 获取列表中的所有图片名称
        selected_files = list(self.listbox2.get(0, tk.END))
        selected_files = sorted(selected_files, key=lambda x: int(re.search(r'\d+', x).group()))
        image_paths = [Load + '\\' + jp for jp in selected_files]
        # 图片名称前增加文字
        for jpg_count in range(len(selected_files)):
            file_names += ['图' + selected_files[jpg_count]]
        # 输出Word文档的路径
        output_path = Save + '\\输出文档.docx' 
#        add_images_to_word(image_paths, file_names, output_path)
        # 创建一个Word文档
        doc = Document()
        describe_count = 0  #初始化描述计数器
        
        self.task_length = len(file_names)  #设定总进度条长度
   #线程控制器使能
        self.is_running = True
        # 开始图片导入文档
        try:
            for img_path, caption in zip(image_paths, file_names):
                if not self.is_running:
                    break
                percent = (describe_count + 1) * 100 // len(file_names)
                # 打开图片并调整大小（这里可以根据需要调整大小）
                img = Image.open(img_path)
                img.thumbnail((Inches(9), Inches(12)))  # 例如，将图片调整为6x6英寸
                img.save(img_path)  # 保存图片（如果需要调整大小）
                if self.add_info.get() and describe_count % self.gap_info.get() == 0:  #间隔42个图片写一段描述
                    imformation = re.split(r'[\s_]+', file_names[describe_count + self.gap_info.get() - 1])
#                    describe = doc.add_paragraph("网压" + imformation[2] + "，" + imformation[5] + "（" + imformation[1] + "），" + imformation[3] + "，" + imformation[4] + "载荷下，" + self.information_entry.get(1.0, tk.END) + "如图" + str(describe_count+1) + "~" + "图" + str(describe_count+self.gap_info.get()) + "所示：")
                    describe = doc.add_paragraph("网压" + imformation[2] + "，" + imformation[5] + "（" + imformation[1] + "），" + imformation[3] + "，" + imformation[4] + "载荷下，" + self.information_entry.get(1.0, tk.END) + "如图" + str(int(imformation[0].replace("图", ""))-self.gap_info.get()+1) + "~" + "图" + str(int(imformation[0].replace("图", ""))) + "所示：")
                    describe.style.font.name = '宋体'  # 设置字体为宋体
                    describe.style.font.size = Pt(12)
#                    describe.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置居中
                # 将图片添加到Word文档中
                doc.add_picture(img_path, width=Inches(6))  # 插入图片，可以指定宽度
                
                # 添加图名（图名段落）
                run = doc.add_paragraph(caption)
        #        run.add_text(caption)
                run.style.font.name = '宋体'  # 设置字体为宋体
                run.style.font.size = Pt(12)  # 设置字号为小四（12pt）实际上应为12磅，但Word中通常小四约
                run.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 设置居中
                if self.add_info.get() and (describe_count + 1) % self.gap_info.get() == 0:  #间隔42个图片写一段描述
                    doc.add_page_break()
                # 保存Word文档
#                doc.save(output_path)
                describe_count += 1
                # 更新进度
                self.queue.put(('progress', percent))
                
            if not selected_files:
                messagebox.showwarning(title = '提示', message = '未选择图片项目')
                self.is_running = False
                self.reset_buttons()
                self.after(100, self.top.destroy)
            # 任务完成
#            if not self.stop_event.is_set():
            if self.is_running:
                self.queue.put(('status', "任务完成!"))
                self.reset_buttons()
                self.after(100, self.top.destroy)
                doc.save(output_path)
        except Exception as e:
            self.queue.put(('error', str(e)))
            self.after(100, self.top.destroy)

    def stop_task(self):
        """停止任务"""
#        self.stop_event.set()
        self.is_running = False
        self.queue.put(('status', "任务已停止"))
        self.reset_buttons()
        self.after(100, self.top.destroy)
#        self.stop_event.clear()        
    def process_queue(self):
        """处理来自工作线程的消息"""
        try:
            while self.is_running:
                msg = self.queue.get_nowait()
#                print(msg[1])
                if msg[0] == 'progress':
                    self.progress_bar['value'] = msg[1]
                    self.progresslabel.config(text=f"进度: {msg[1]}%")
                elif msg[0] == 'status':
                    messagebox.showinfo(title = '提示', message = msg[1])
#                    self.top.destroy()
                    self.reset_buttons()
                
                elif msg[0] == 'error':
                    messagebox.showerror(title = '提示', message = f'发生了一个错误：{msg[1]}，导出文档失败')
#                    self.top.destroy()
                    self.reset_buttons()
            msg = self.queue.get_nowait()
            if msg[0] == 'status':
                messagebox.showinfo(title = '提示', message = msg[1])
                
#            print(msg[1])
        except queue.Empty:
            pass
        
        # 继续定期检查队列
        self.after(100, self.process_queue)
    
    def reset_buttons(self):
        """重置按钮状态"""
        self.run_button2.config(state=tk.NORMAL)
#        self.stop_button.config(state=tk.DISABLED)
    
#    def on_close(self):
#        """窗口关闭事件处理"""
#        self.after(100, self.top.destroy)
#        self.stop_task()
        
#        self.top.destroy()
#        print(self.queue.get_nowait())
