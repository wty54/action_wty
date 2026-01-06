# -*- coding: utf-8 -*-
"""
Created on Thu Dec 11 14:45:59 2025

@author: thinkpad
"""
from docx import Document


def smart_merge_with_page_breaks(docA_path, docB_path, output_path, a_pages=3, b_pages=2):
    """
    智能合并：使用分页符控制页面
    """
    
    docA = Document(docA_path)
    docB = Document(docB_path)
    merged_doc = Document()
    
    # 获取两个文档的所有元素（段落和表格）
    a_elements = []
    b_elements = []
    
    # 收集文档A的所有段落和表格
    for element in docA.element.body:
        a_elements.append(element)
    
    # 收集文档B的所有段落和表格
    for element in docB.element.body:
        b_elements.append(element)
    
    # 计算要处理的组数
    total_a_groups = len(a_elements) // (a_pages * 10) + 1  # 假设10个元素一页
    total_b_groups = len(b_elements) // (b_pages * 10) + 1
    total_groups = min(total_a_groups, total_b_groups)
    
    print(f"预计合并 {total_groups} 组内容")
    
    # 合并逻辑
    for group in range(total_groups):
        print(f"处理第 {group + 1} 组...")
        
        # 添加A文档内容
        a_start = group * a_pages * 10
        a_end = min(a_start + a_pages * 10, len(a_elements))
        
        for i in range(a_start, a_end):
            merged_doc.element.body.append(a_elements[i])
        
        # 添加分页符（如果是最后一组且后面没有B内容，则不添加）
        if group < total_groups - 1 or (group == total_groups - 1 and b_pages > 0):
            merged_doc.add_page_break()
        
        # 添加B文档内容
        b_start = group * b_pages * 10
        b_end = min(b_start + b_pages * 10, len(b_elements))
        
        for i in range(b_start, b_end):
            merged_doc.element.body.append(b_elements[i])
        
        # 如果不是最后一组，添加分页符
        if group < total_groups - 1:
            merged_doc.add_page_break()
    
    # 保存文档
    merged_doc.save(output_path)
    print(f"合并完成！文档已保存到: {output_path}")
    
    return output_path

# 使用示例
docA_path = "C:\\Users\\thinkpad\\Desktop\\20251210\\嘉兴南_枫南_基本数据表.docx"
docB_path = "C:\\Users\\thinkpad\\Desktop\\20251210\\嘉兴南_枫南.docx"
output_path = "C:\\Users\\thinkpad\\Desktop\\20251210\\合并文档.docx"
smart_merge_with_page_breaks(docA_path, docB_path, output_path, a_pages=8, b_pages=4)