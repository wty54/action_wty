# -*- coding: utf-8 -*-
"""
Created on Wed Dec 10 16:39:30 2025

@author: thinkpad
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from main import sorted_files
import pandas as pd
import os
import re

# 筛选路径下的.xls文件
def collect_excel(root_dir):
    file_list = []
    for filename in os.listdir(root_dir):
        if filename.endswith('.xls'):
#            file_path = os.path.join(root_dir, filename)
            file_list += [filename]
    file_list = sorted_files(file_list, "25000V", "基本")
    return file_list


def add_border_to_table(table):
    """
    为表格添加边框
    """
    tbl = table._tbl  # 获取表格的XML元素
    tblPr = tbl.tblPr
    
    # 设置表格边框
    tblBorders = OxmlElement('w:tblBorders')
    
    # 设置上边框
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    tblBorders.append(top)
    
    # 设置下边框
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)
    
    # 设置左边框
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '4')
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), '000000')
    tblBorders.append(left)
    
    # 设置右边框
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '4')
    right.set(qn('w:space'), '0')
    right.set(qn('w:color'), '000000')
    tblBorders.append(right)
    
    # 设置内部水平边框
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), '000000')
    tblBorders.append(insideH)
    
    # 设置内部垂直边框
    insideV = OxmlElement('w:insideV')
    insideV.set(qn('w:val'), 'single')
    insideV.set(qn('w:sz'), '4')
    insideV.set(qn('w:space'), '0')
    insideV.set(qn('w:color'), '000000')
    tblBorders.append(insideV)
    
    tblPr.append(tblBorders)
    # ========== 设置表格字体为宋体小五 ==========
def set_cell_font(cell, font_name='宋体', font_size=Pt(9), bold=False):
    """设置单元格字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
def excel_to_word_pandas(file_path, word_path, excel_num,
                                max_rows=23, max_cols=16,
                                sheet_name="运行基本数据表"):
    #  初始化变量
    data_kilo = 0  #总里程    
    data_average_time = 0  #总旅行时间
    # 创建Word文档
    doc = Document()
    
    # ========== 设置页面为横向 ==========
    section = doc.sections[0]
#    section.orientation = WD_ORIENT.LANDSCAPE  # 设置为横向
    section.page_width = Cm(29.7)  # A4横向的宽度（29.7cm）
    section.page_height = Cm(21.0)  # A4横向的高度（21.0cm）
    # 设置宋体五号
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal'].font.size = Pt(12)
    # 设置页边距（单位：厘米）
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
#    # 添加标题
#    doc.add_heading('Excel数据提取', level=1)
    file_list = collect_excel(file_path)
    for file in file_list:
        information = re.split(r'[_.]', file)
        excel_path = file_path + "\\" + file 
        try:
            
            df = pd.read_excel(
                excel_path,
                nrows=max_rows,
                header=0
            )
            
            # 截取前16列
            if df.shape[1] > max_cols:
                df = df.iloc[:, :max_cols]
            print(f"成功读取数据: {df.shape[0]}行 x {df.shape[1]}列，文件名为: {file}")
            data_kilo += df.iloc[5, 6]
            data_average_time += df.iloc[7, 6]
            data_energy_consume = df.iloc[12, 6]
            data_energy_regenerate = df.iloc[13, 6]
            data_average_velocity = df.iloc[9, 6]
            
            # 计算100%、50%、0%再生率的能耗
            consume_50 = round(df.iloc[:, 11] - df.iloc[:, 12] / 2.0, 2)
            df.insert(14, '区间总能耗(再生50%)(kWh)', consume_50)
            consume_0 = df.iloc[:, 11]
            df.insert(15, '区间总能耗(无再生)(kWh)', consume_0)
            consume_kilo_50 = round(df.iloc[:, 14] / df.iloc[:, 6], 2)
            df.insert(17, '区间平均公里耗电量(再生50%)(kWh/km)', consume_kilo_50)
            consume_kilo_0 = round(df.iloc[:, 15] / df.iloc[:, 6], 2)
            df.insert(18, '区间平均公里耗电量(无再生)(kWh/km)', consume_kilo_0)
            consume_ratio_50 = round(df.iloc[:, 19] / 2.0, 2)
            df.insert(20, '再生制动率(再生50%)', consume_ratio_50)
            consume_ratio_0 = round(df.iloc[:, 19] * 0)
            df.insert(21, '再生制动率(无再生)', consume_ratio_0)
#            print(value)
        except Exception as e:
            print(f"读取Excel文件失败: {e}")
            return
        doc.add_paragraph(f"列车在网压{information[1]}、{information[2]}、{information[3]}载荷、节能运行模式（惰行）条件下的{information[4]}仿真基本运行数据见表{excel_num}，其中累计能耗为{data_energy_consume}kW·h，再生能量为{data_energy_regenerate}kW·h，平均旅行速度为{data_average_velocity}km/h。换行").alignment = WD_ALIGN_PARAGRAPH.LEFT
        #  添加标题
        doc.add_paragraph(f"表{excel_num}空格{information[1]}_{information[2]}_{information[3]}_基本运行数据({information[4]})").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 创建数据表格（增加一行用于表头）
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        # 不设置table.style，手动添加边框
        add_border_to_table(table)
        
        # 设置列宽
        for col in table.columns:
            col.width = Inches(1.2)
        
        # 添加表头
        header_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            header_cells[i].text = str(col_name)
            set_cell_font(header_cells[i], '宋体', Pt(9), bold=False)
        # 填充数据
        for i, row in df.iterrows():
            row_cells = table.rows[i + 1].cells
            for j, value in enumerate(row):
                # 处理NaN值
                if pd.isna(value):
                    cell_text = ''
                else:
                    cell_text = str(value)
                row_cells[j].text = cell_text
                set_cell_font(row_cells[j], '宋体', Pt(9), bold=False)
        # 水平合并
        for row_idx in range(6, 23):  # 7对应第8行，23对应第24行
            base_cell = table.rows[row_idx].cells[2]  # 第3列
            for col_idx in range(3, 23):  # 从第4列到第23列
                if col_idx < len(table.columns):
                    cell_to_merge = table.rows[row_idx].cells[col_idx]
                    base_cell.merge(cell_to_merge)
        if excel_num % 2 == 0:
            data_total_velocity = round(data_kilo*3600/(data_average_time+180), 2)  # 计算总旅行速度
            doc.add_paragraph(f"考虑折返时间180s，计算全线平均旅行速度为{data_total_velocity}km/h。")
            data_kilo = 0
            data_average_time = 0
        doc.add_section()  #插入下一页
        excel_num += 1
    # 设置合并后单元格
#    base_cell.text = merge_content
#    base_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中
    # 保存文档
    doc.save(word_path)
    print(f"文档已保存: {word_path}")
#    return doc
# 使用示例
if __name__ == "__main__":
    file_path = "C:\\Users\\thinkpad\\Desktop\\20251212\\test"
    word_path = "C:\\Users\\thinkpad\\Desktop\\20251212\\基本数据表.docx"
#    file_path = "C:\\Users\\thinkpad\\Desktop\\test"
#    word_path = "C:\\Users\\thinkpad\\Desktop\\test\\output.docx"
    excel_num = 11  #初始表格编号
#    file_list = collect_excel(file_path)  #用于存储.xls文件名
    excel_to_word_pandas(file_path, word_path, excel_num)
#    doc.save(word_path)