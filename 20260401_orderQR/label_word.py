# -*- coding: utf-8 -*-
"""
Created on Tue Mar 24 11:18:51 2026

@author: thinkpad
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
import pandas as pd
import qrcode
from PIL import Image
from io import BytesIO
import os
import re
from docxcompose.composer import Composer
import sys

def replacement_construct(replacements, data_catagory, count, QR_string):
    r = replacements[count]
#    date = re.sub(r'(\d{4})年(\d{1,2})月(\d{1,2})日', lambda m : f"{m.group(1)[-2:]}{m.group(2).zfill(2)}{m.group(3).zfill(2)}", r[9])
    replacement = {}
    for i, catagory in enumerate(data_catagory):
        if catagory == "供应商公司":
            replacement["<" + catagory + ">"] = str(r[i])
        else:
            replacement["<" + catagory + ">"] = catagory + "：" + str(r[i])
#    replacement = {
#            "<项目名称>" : "项目名称：" + str(r[1]),
#            "<采购订单>" : "采购订单：" + str(r[2]),
#            "<物资编号>" : "物资编号：" + str(r[3]),
#            "<规格型号>" : "规格型号：" + str(r[4]),
#            "<物料名称>" : "物料名称：" + str(r[5]),
#            "<计量单位>" : "计量单位：" + str(r[6]),
#            "<数量>" : "数量：" + str(r[7]),
#            "<序列号>" : "序列号：" + str(r[8]),
#            "<生产日期>" : "生产日期：" + str(r[9]),
#            "<保质期>" : "保质期：" + str(r[10]),
#            "<列数>" : "列数：" + str(r[11]),
#            "<供应商代码>" : "供应商代码：" + str(r[12]),
#            "<供应商公司>" : str(r[13]),
#            "<备用1>" : "备用1：" + str(r[14]),
#            "<备用2>" : "备用2：" + str(r[15]),
#            "<备用3>" : "备用3：" + str(r[16]),
#            "<备用4>" : "备用4：" + str(r[17]),
#            "<备用5>" : "备用5：" + str(r[18]),
##            "<二维码>" : r[3] + ";" + r[12] + ";" + r[8] + ";" + date + ";;;;;",
#            }
    replacement["<二维码>"] = QR_info_transform(QR_string, replacement)
    return replacement

def replace_information(doc_path, output_path, replacements, data_catagory, QR_info):
    """
    替换内容并保留原有格式
    """
    doc = Document(doc_path)
    r_count = 1
    for table in doc.tables:
        replacement = replacement_construct(replacements, data_catagory, r_count, QR_info["information"])
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text
                # 检查是否需要替换
                for search_text, replace_text in replacement.items():
                    if search_text in cell_text:
                        cell_text2 = cell_text.replace(search_text, replace_text)
                        # 清空单元格
                        for paragraph in cell.paragraphs:
                            paragraph.clear()
                        if search_text == "<二维码>":
                            run = cell.paragraphs[0].add_run()
                            run.add_picture(QRcode_producer(cell_text2), width=Cm(QR_info["size"]), height=Cm(QR_info["size"]))
                        # 添加新文本
                        else:
                            cell.paragraphs[0].add_run(cell_text2)
                            set_cell_font(cell, '宋体', Pt(16), bold=True)
#                        print(f"已替换：'{search_text}'")
#                        break  # 一个单元格仅替换一次
        r_count += 1                  
    doc.save(output_path)
#    print(f"处理完成！保存到: {output_path}")

def set_cell_font(cell, font_name='宋体', font_size=Pt(16), bold=True):
    """设置单元格字体"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size
            run.bold = bold

def QRcode_producer(QR_information):
    # 创建二维码实例
    qr = qrcode.QRCode(
        version=1,  # 二维码的版本，范围是1到40
        error_correction=qrcode.constants.ERROR_CORRECT_L,  # 错误更正级别，'L', 'M', 'Q', 'H'
        box_size=10,  # 每个小方块代表的像素大小
        border=1,  # 边框的大小
    )
    
    # 添加数据到二维码中
    qr.add_data(QR_information)
    qr.make(fit=True)  # 生成二维码

    # 创建二维码图像
    img = qr.make_image(fill_color="black", back_color="white").convert("RGB")
    # 将PIL Image对象保存到内存流中
    img_stream = BytesIO()
    img.save(img_stream, format='PNG')
    img_stream.seek(0)  # 重置流位置到开头
    
    return img_stream

def copy_tables(model_path, f_count):
    # 创建临时文件
    if f_count <= 0:
        return model_path
    master = Document(model_path)
    composer = Composer(master)
        # 创建 10 个临时副本
    for i in range(1, f_count):
        temp_name = f"temp_copy_{i}.docx"
#            temp_files.append(temp_name)
        
        # 复制原文件
        with open(model_path, "rb") as src:
            with open(temp_name, "wb") as dst:
                dst.write(src.read())
    
    # 合并所有临时文件
        doc = Document(temp_name)
        composer.append(doc)
        
        # 保存最终文档
#            composer.save(output_path)
        composer_stream = BytesIO()
        composer.save(composer_stream)
        composer_stream.seek(0)
#        print(f"合成文件: temp_copy_{i}.docx")
        if os.path.exists(temp_name):
            os.remove(temp_name)
#            print(f"删除缓存文件: temp_copy_{i}.docx")
    return composer_stream
    
def excel_model(excel_path):
    try:
        df = pd.read_excel(excel_path, dtype = str, header=0)
        row_count = df.shape[0]
        data_catagory = df.columns.tolist()
        data_list = df.values.tolist()
        replacements = {i + 1 : row for i, row in enumerate(data_list)}
#        print(replacements)
        return replacements, row_count, data_catagory
    except Exception as e:
#        print(f"读取Excel文件失败: {e}")
        return

def QR_info_transform(QR_string, replacement):
    QR_list = QR_string.split(';')
    for i, QR in enumerate(QR_list):
        if QR:
            string = replacement["<" + QR + ">"]
            match = re.search(rf'{QR}：(.*)', string)
            if match:
                string = match.group(1)
                if QR == "生产日期":
                    string = re.sub(r'(\d{4})年(\d{1,2})月(\d{1,2})日', lambda m : f"{m.group(1)[-2:]}{m.group(2).zfill(2)}{m.group(3).zfill(2)}", string)
                QR_list[i] = string
    QR_info = ';'.join(QR_list)
#    print(QR_out)
    return QR_info

def get_resource_path(relative_path):
    """获取资源的绝对路径，兼容开发环境和打包后的exe"""
    if getattr(sys, 'frozen', False):
        # 打包后的exe运行
        base_path = sys._MEIPASS
    else:
        # 开发环境运行
        base_path = os.path.abspath(".")
    
    return os.path.join(base_path, relative_path)


def start(model1_path, output1_path, excel_path, QR_info1):
    replacements, row_count, data_catagory = excel_model(excel_path)
    composer_stream1 = copy_tables(model1_path, row_count)
    replace_information(composer_stream1, output1_path, replacements, data_catagory, QR_info1)

if __name__ == "__main__":
    model1_path = "C:\\Users\\thinkpad\\Desktop\\test\\唛头模板.docx"
    model2_path = "C:\\Users\\thinkpad\\Desktop\\test\\标签模板.docx"
    output1_path = "C:\\Users\\thinkpad\\Desktop\\test\\output1.docx"
    output2_path = "C:\\Users\\thinkpad\\Desktop\\test\\output2.docx"
    excel_path = "C:\\Users\\thinkpad\\Desktop\\test\\发货信息模板.xlsx"
    QR_info1 = { "information": "物资编号;供应商代码;序列号;生产日期;;;;;",
                "size" : 3,
                }
    QR_info2 = { "information": "物资编号;供应商代码;序列号;生产日期;;;;;",
                "size" : 3.8,
                }
    start(model1_path, output1_path, excel_path, QR_info1)
#    composer_stream2 = copy_tables(model2_path, row_count)
#    replace_information(composer_stream2, output2_path, replacements, QR_info2)

    
    
    